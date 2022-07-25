import argparse
from docx import Document

def main(args):
    document = Document(args.docx)
    frontMatter = list()
    backMatter = list()
    # Skip first line, which should be title
    for i in document.paragraphs[1:]:
        para = document.paragraphs[i].text
        try:
            int(para[0])
        except TypeError:
            continue


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument('docx',
                        help='Path of the docx file',
                        type=str)
    args = parser.parse_args()
    main(args)
